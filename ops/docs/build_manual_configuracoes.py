from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\projetos\AVMD_System")
OUT = ROOT / "Manual_de_Configuracoes_AVMD_CertiID.docx"
LOGO = ROOT / "public" / "logo-certiid.png"

IMAGES = {
    "usuarios": Path(r"C:\Users\manto\AppData\Local\Temp\codex-clipboard-05049845-73af-42d0-9c26-63817459d807.png"),
    "parceiros": Path(r"C:\Users\manto\AppData\Local\Temp\codex-clipboard-cee84127-2cd9-4e9d-a77b-c45d71d68472.png"),
    "modelo": Path(r"C:\Users\manto\AppData\Local\Temp\codex-clipboard-85325707-7cba-4b45-90c8-351c70b2d8f4.png"),
    "faixas": Path(r"C:\Users\manto\AppData\Local\Temp\codex-clipboard-a08c6f1d-214a-4899-8402-3397de5bf5a8.png"),
    "remuneracao": Path(r"C:\Users\manto\AppData\Local\Temp\codex-clipboard-af1da87e-fcaa-44b8-a6a0-c89aa0e878e5.png"),
}

BLUE = "155EEF"
NAVY = "17335C"
ORANGE = "F28C28"
LIGHT = "E8EEF5"
PALE = "F5F7FA"
GRAY = "5C677D"
RED = "A61B1B"
CURRENT_STEP = 0


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_font(run, size=11, bold=False, color="222222", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, 30, True, NAVY)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(16)
        r2 = p2.add_run(subtitle)
        set_font(r2, 14, False, GRAY)


def h1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.keep_with_next = True
    return p


def h2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.keep_with_next = True
    return p


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, 11, True, NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.38)
    p.paragraph_format.first_line_indent = Inches(-.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_font(p.add_run(text))
    return p


def reset_steps(doc):
    global CURRENT_STEP
    CURRENT_STEP = 0


def step(doc, text):
    global CURRENT_STEP
    CURRENT_STEP += 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.18)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    set_font(p.add_run(f"Etapa {CURRENT_STEP} — "), 11, True, NAVY)
    set_font(p.add_run(text))
    return p


def note(doc, title, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.3)
    set_cell_fill(cell, "EEF4FF" if color == BLUE else "FFF2F2")
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, 10.5, True, color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_repeat_table_header(t.rows[0])
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_fill(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_font(p.add_run(header), 10, True, NAVY)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def screenshot(doc, key, caption, width=6.25):
    path = IMAGES[key]
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    set_font(p2.add_run(caption), 9, False, GRAY, True)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 9),
    ("Heading 2", 13, BLUE, 13, 6),
    ("Heading 3", 11.5, NAVY, 9, 4),
):
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("AVMD SYSTEM | Manual de Configurações"), 8.5, False, GRAY)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(fp.add_run("CertiID • versão operacional 1.0 • página "), 8.5, False, GRAY)
add_field(fp, "PAGE")

if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.add_run().add_picture(str(LOGO), width=Inches(1.9))
add_title(doc, "Manual de Configurações", "AVMD SYSTEM • CRM CertiID")
body(doc, "Guia administrativo para implantação, manutenção e auditoria das configurações operacionais, comerciais, fiscais e de atendimento.")
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(36)
set_font(meta.add_run("Edição: julho de 2026\nAmbiente: Produção\nPúblico: administradores e supervisores"), 10.5, False, GRAY)
doc.add_page_break()

h1(doc, "Sumário operacional")
table(doc, ["Capítulo", "Conteúdo"], [
    ("1", "Princípios de configuração e acesso"),
    ("2", "Usuários, perfis, vínculos e permissões"),
    ("3", "Parceiros e agentes permitidos"),
    ("4", "Pontos de atendimento e hierarquia"),
    ("5", "Modelo comercial, imposto e RBT12"),
    ("6", "Remuneração, comissões e repasses"),
    ("7", "Tabelas de preço e descontos"),
    ("8", "Fluxos de venda, validação e representante legal"),
    ("9", "Chat, filas, supervisão e histórico"),
    ("10", "Fiscal/NFS-e, relatórios, segurança e checklist"),
], [0.8, 5.5])
note(doc, "Como usar este manual", "Siga primeiro o checklist de implantação. Depois consulte o capítulo específico sempre que alterar uma regra. Mudanças de imposto, preço ou comissão devem ser registradas com data de vigência.")

h1(doc, "1. Princípios de configuração")
body(doc, "O sistema separa cadastro, acesso, operação e remuneração. Um mesmo participante pode acumular papéis, mas cada papel deve ser marcado no cadastro para aparecer nas listas compatíveis.")
table(doc, ["Camada", "O que controla", "Exemplo"], [
    ("Cadastro", "Identidade e dados do participante", "Parceiro Rogerio Ribeiro"),
    ("Acesso", "O que o usuário pode visualizar e executar", "Administrador ou Agente de Registro"),
    ("Vínculo", "Em nome de quem o usuário opera", "Parceiro Vendedor"),
    ("Ponto", "Local operacional e hierarquia", "CertiID Matriz"),
    ("Modelo comercial", "Imposto, retenção e divisão financeira", "Integrado ou Revenda"),
], [1.25, 2.75, 2.3])
note(doc, "Regra de segurança", "Não compartilhe logins. Cada pessoa deve possuir usuário próprio. A senha não deve ser armazenada no cadastro do parceiro.", RED)

h1(doc, "2. Usuários, perfis e vínculos")
h2(doc, "2.1 Caminho")
body(doc, "Configurações > Usuários > selecione o usuário > Editar.")
h2(doc, "2.2 Perfil de acesso")
table(doc, ["Perfil", "Uso recomendado"], [
    ("Administrador", "Configuração global e acesso amplo."),
    ("Supervisor do Chat", "Visualiza e distribui todas as conversas."),
    ("Agente de Registro", "Validações, agenda e operação autorizada."),
    ("Parceiro Vendedor", "Vendas e clientes vinculados."),
    ("Usuário", "Acesso limitado por permissões marcadas."),
], [2.0, 4.3])
h2(doc, "2.3 Vínculo do usuário")
table(doc, ["Vínculo", "Quando usar"], [
    ("Agente de Registro", "Pessoa habilitada a realizar validações."),
    ("Parceiro Vendedor", "Vende por acesso ou link próprio."),
    ("Parceiro Contador", "Indica ou vende para sua carteira."),
    ("Usuário comum", "Colaborador interno sem vínculo comercial."),
    ("Cliente do portal", "Cliente externo com acesso ao próprio portal."),
], [2.0, 4.3])
reset_steps(doc)
step(doc, "Selecione o perfil de acesso.")
step(doc, "Selecione o vínculo funcional.")
step(doc, "No campo Nome do vínculo, escolha um cadastro compatível.")
step(doc, "Marque somente os módulos necessários.")
step(doc, "Salve e teste o acesso em sessão separada.")
screenshot(doc, "usuarios", "Tela de edição do usuário e escolha do vínculo funcional.")
note(doc, "Importante", "A opção genérica Parceiro foi retirada. Utilize Parceiro Vendedor ou Parceiro Contador para evitar ambiguidade.")

h1(doc, "3. Parceiros")
h2(doc, "3.1 Cadastro")
body(doc, "Parceiros > Novo Parceiro ou ícone de edição. Informe documento, razão social/nome, contato, endereço, IBGE, tipo, papéis adicionais e dados bancários quando houver repasse.")
bullet(doc, "O cadastro é único: não recrie a mesma pessoa para cada papel.")
bullet(doc, "Marque os papéis adicionais: vendedor, agente de registro e/ou contador.")
bullet(doc, "O local de atendimento deve ser escolhido na lista de pontos cadastrados.")
bullet(doc, "Dados bancários são necessários somente quando houver pagamento.")
h2(doc, "3.2 Agentes permitidos")
body(doc, "A seção Agentes de Registro Permitidos limita quem pode atender vendas e validações daquele parceiro. Ela não define comissão.")
reset_steps(doc)
step(doc, "Escolha o agente.")
step(doc, "Defina o ponto preferencial ou deixe Sem ponto fixo.")
step(doc, "Clique em Vincular agente.")
step(doc, "Use inativar para suspender temporariamente; excluir remove o vínculo.")
screenshot(doc, "parceiros", "Configuração do parceiro e agentes de registro permitidos.")

h1(doc, "4. Pontos de atendimento e hierarquia")
body(doc, "Configurações > Pontos de Atendimento. Cada ponto deve ter nome, código, cidade, UF e status. O botão Hierarquia abre os agentes e participantes vinculados.")
reset_steps(doc)
step(doc, "Abra Hierarquia no ponto desejado.")
step(doc, "Vincule o agente raiz do ponto.")
step(doc, "Adicione vendedores ou participantes subordinados quando necessário.")
step(doc, "Defina se o vendedor pode usar todos os agentes ou apenas um agente específico.")
step(doc, "Abra Modelo e comissões no participante que lidera a operação.")
table(doc, ["Ponto", "Agentes atualmente previstos"], [
    ("CertiID Matriz", "Renata Mantovan e Daniel Morgado Mantovan"),
    ("CertiID BH", "Isabella Vidal"),
    ("CertiID Mauá", "Ingrid Braz"),
    ("CertiID SJC", "Alice Maciel"),
], [2.0, 4.3])
note(doc, "Diferença essencial", "Agente permitido no parceiro controla atendimento. Agente vinculado ao ponto controla a hierarquia operacional. São relações diferentes.")

h1(doc, "5. Modelo comercial e imposto")
h2(doc, "5.1 Caminho")
body(doc, "Configurações > Pontos de Atendimento > Hierarquia > participante > Modelo e comissões.")
h2(doc, "5.2 Integrado")
body(doc, "A CertiID define o preço de venda. Todas as comissões usam o valor da venda após o imposto. O saldo não distribuído permanece com a operação integrada.")
h2(doc, "5.3 Revenda")
body(doc, "A empresa trava a retenção/custo base. Do valor após imposto são descontados a retenção e os repasses. O saldo pertence ao revendedor.")
h2(doc, "5.4 Simples Nacional — Anexo III")
body(doc, "Selecione Simples — Anexo III e informe o RBT12. O sistema identifica a faixa e calcula a alíquota efetiva pela fórmula: (RBT12 × alíquota nominal − parcela a deduzir) ÷ RBT12.")
table(doc, ["RBT12 informado", "Faixa", "Alíquota efetiva"], [
    ("R$ 274.489,49", "2ª faixa", "7,7900%"),
], [2.2, 1.6, 2.5])
body(doc, "Exemplo de venda: R$ 200,00 − R$ 15,58 de imposto = R$ 184,42 de base para remunerações.")
screenshot(doc, "modelo", "Modelo Integrado com cálculo automático pelo Anexo III.")
note(doc, "Atualização mensal", "Atualize o RBT12 após o fechamento de cada competência. O sistema calcula a alíquota; a origem contábil do RBT12 deve ser conferida com a contabilidade.", RED)

h1(doc, "6. Remuneração e repasses")
h2(doc, "6.1 Regra vigente")
body(doc, "A regra operacional efetiva está em Remuneração do Agente. Ela permite valor fixo ou percentual, por ponto e por tipo de documento.")
bullet(doc, "Existe uma única comissão de validação por atendimento.")
bullet(doc, "Podem existir vários repasses de venda.")
bullet(doc, "Percentuais são aplicados sobre a venda após o imposto.")
bullet(doc, "O sistema bloqueia estruturas cujo total ultrapasse a base disponível.")
h2(doc, "6.2 Remuneração em cascata")
reset_steps(doc)
step(doc, "Escolha Quem recebe.")
step(doc, "Mantenha a base Venda após imposto.")
step(doc, "Selecione Percentual ou Valor fixo.")
step(doc, "Informe o valor e adicione o repasse.")
step(doc, "Repita para contador, vendedor ou intermediário, se aplicável.")
table(doc, ["Venda", "Imposto 7,79%", "Base", "AGR 20%", "Contador 10%", "Saldo"], [
    ("R$ 200,00", "R$ 15,58", "R$ 184,42", "R$ 36,88", "R$ 18,44", "R$ 129,10"),
], [1.0, 1.1, 1.0, 1.0, 1.1, 1.1])
h2(doc, "6.3 Faixas de comissão")
body(doc, "A tela Faixas de Comissão existe como cadastro de faixas por volume, mas ainda não é a fonte do cálculo financeiro. Não configure valores conflitantes nela.")
screenshot(doc, "faixas", "Faixas por volume: referência atual, ainda não ligada ao cálculo efetivo.", 5.6)
screenshot(doc, "remuneracao", "Remuneração do agente: regra atualmente utilizada no cálculo.", 5.6)
note(doc, "Evite duplicidade", "Até a unificação das telas, utilize Remuneração do Agente como regra oficial de pagamento.", RED)

h1(doc, "7. Tabelas de preço e descontos")
body(doc, "Tabelas de preço podem ser vinculadas a parceiros, vendedores e contadores. Um participante pode acessar mais de uma tabela.")
bullet(doc, "No modelo Integrado, o preço de venda é controlado pela operação.")
bullet(doc, "No modelo Revenda, configure o preço base/retenção por produto.")
bullet(doc, "A venda não pode ficar abaixo do mínimo definido.")
bullet(doc, "Descontos devem ser autorizados e registrados antes do pagamento.")
note(doc, "Tela de venda", "A aba operacional de vendas mostra apenas o valor da venda. Imposto, custos e repasses devem aparecer nos relatórios financeiros.")

h1(doc, "8. Venda, validação e representante legal")
h2(doc, "8.1 Venda")
reset_steps(doc)
step(doc, "Selecione o cliente ou conclua o cadastro obrigatório.")
step(doc, "Escolha produto, tabela, vendedor/parceiro, ponto e forma de pagamento.")
step(doc, "O sistema registra a estrutura comercial como fotografia da venda.")
step(doc, "Após emissão, alterações de configuração não devem modificar vendas antigas.")
h2(doc, "8.2 Validação")
body(doc, "O agente efetivamente responsável deve ser registrado no pedido. A comissão de validação é calculada uma única vez sobre a base após imposto.")
h2(doc, "8.3 Representante legal")
body(doc, "Para e-CNPJ, valide o representante legal antes de prosseguir. O CPF completo deve ser informado e armazenado com proteção adequada. Para e-PJ e NF-e, a regra pode ser dispensada conforme o produto.")
note(doc, "LGPD", "CPF completo é dado pessoal. Restrinja acesso, registre finalidade e mantenha trilha de auditoria.", RED)

h1(doc, "9. Chat, filas e supervisão")
bullet(doc, "Administrador e Supervisor do Chat podem visualizar todas as conversas.")
bullet(doc, "Atendente visualiza conversas atribuídas ou permitidas pela fila.")
bullet(doc, "Arquivar remove da visão operacional, mas preserva cliente e histórico.")
bullet(doc, "Nova mensagem recebida deve reabrir ou reativar a conversa.")
bullet(doc, "Excluir não deve ser usado como arquivamento.")
bullet(doc, "Áudio, imagem e arquivo precisam permanecer acessíveis pelo histórico.")
reset_steps(doc)
step(doc, "Defina o modo: IA, humano ou IA e humano.")
step(doc, "Atribua a conversa ao atendente responsável.")
step(doc, "Salve o contato quando houver dados suficientes.")
step(doc, "Ao concluir, arquive com histórico e resumo.")
note(doc, "Falha de envio", "Se uma mensagem falhar, preserve o texto digitado, mostre o erro e permita reenviar. Não altere a fila antes de confirmar o envio.", RED)

h1(doc, "10. Fiscal, relatórios e auditoria")
h2(doc, "10.1 NFS-e")
body(doc, "Configurações > Fiscal/NFS-e. A configuração é exclusiva da Certifast Certificação Digital Ltda, CNPJ 20.776.537/0001-55, município de São Bernardo do Campo.")
bullet(doc, "Serviço: 17.02/102818/1241.")
bullet(doc, "CNAE: 8299-7/99.")
bullet(doc, "Ambiente, credenciais e certificado devem ser protegidos.")
bullet(doc, "Antes de emitir, valide tomador, endereço, CEP, e-mail e telefone.")
bullet(doc, "Cancelamento exige justificativa e retorno confirmado pela prefeitura.")
h2(doc, "10.2 Relatórios")
body(doc, "Relatórios operacionais devem filtrar por período, pedido, protocolo, parceiro, vendedor, agente e status. Relatórios de repasse devem destacar venda, imposto, base, validação, venda, custos e saldo.")
bullet(doc, "Exportações disponíveis: XLSX, XLS e CSV, conforme a tela.")
bullet(doc, "Filtros e ordem de colunas devem ser preservados por usuário.")
bullet(doc, "Relatórios salvos devem registrar os filtros utilizados.")
h2(doc, "10.3 Checklist de implantação")
for item in [
    "Usuários possuem perfis e vínculos corretos.",
    "Parceiros não estão duplicados e têm papéis adicionais marcados.",
    "Pontos possuem agentes e hierarquia vinculados.",
    "Modelo Integrado ou Revenda foi salvo em cada operação.",
    "RBT12 e alíquota foram conferidos.",
    "Preço base da Revenda foi cadastrado por produto.",
    "Remunerações não ultrapassam a base após imposto.",
    "Tabelas de preço estão vinculadas aos participantes corretos.",
    "Chat foi testado com texto, áudio, imagem e arquivo.",
    "NFS-e foi testada com emissão, visualização, envio e cancelamento.",
    "Relatórios e exportações foram conferidos.",
]:
    bullet(doc, "☐ " + item)

h1(doc, "Registro de alterações")
table(doc, ["Data", "Responsável", "Configuração alterada", "Vigência"], [
    ("____/____/______", "________________", "________________________________", "____/____/______"),
], [1.2, 1.5, 2.6, 1.2])

doc.core_properties.title = "Manual de Configurações do AVMD SYSTEM"
doc.core_properties.subject = "CRM CertiID - configurações administrativas e operacionais"
doc.core_properties.author = "CertiID"
doc.core_properties.keywords = "AVMD, CertiID, CRM, configurações, comissões, NFS-e, chat"
doc.save(OUT)
print(OUT)
